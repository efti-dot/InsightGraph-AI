import os
from docx import Document
import plotly.io as pio
from docx.shared import Inches
from fpdf import FPDF

REPORTS_ROOT = os.path.join("storage", "reports")

PDF_SAFE_REPLACEMENTS = {
    "\u2014": "-",
    "\u2013": "-", 
    "\u2018": "'",
    "\u2019": "'",
    "\u201c": '"',
    "\u201d": '"',
    "\u2026": "...",
    "\u2022": "-",
}

def _report_dir(project_id: str) -> str:
    path = os.path.join(REPORTS_ROOT, project_id)
    os.makedirs(path, exist_ok=True)
    return path

def sanitize_for_pdf(text: str) -> str:
    for old, new in PDF_SAFE_REPLACEMENTS.items():
        text = text.replace(old, new)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def _render_chart_images(charts: list[dict], out_dir: str) -> list[str]:
    """Renders each chart's Plotly JSON to a PNG (needs the kaleido package)."""
    image_paths = []
    for i, chart in enumerate(charts):
        try:
            fig = pio.from_json(chart["figure_json"])
            image_path = os.path.join(out_dir, f"chart_{i}.png")
            fig.write_image(image_path, width=800, height=450, scale=2)
            image_paths.append(image_path)
        except Exception as exc:
            print(f"[export_service] failed to render chart {i}: {exc}")
    return image_paths


def export_docx(project_id: str, draft_report: str, charts: list[dict]) -> str:
    out_dir = _report_dir(project_id)
    doc = Document()

    for line in draft_report.splitlines():
        stripped = line.strip()
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped:
            doc.add_paragraph(stripped)

    if charts:
        doc.add_heading("Chart Images", level=2)
        for img_path in _render_chart_images(charts, out_dir):
            doc.add_picture(img_path, width=Inches(6))

    path = os.path.join(out_dir, "report.docx")
    doc.save(path)
    return path


def export_pdf(project_id: str, draft_report: str, charts: list[dict]) -> str:
    out_dir = _report_dir(project_id)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)

    for line in draft_report.splitlines():
        stripped = sanitize_for_pdf(line.strip())
        pdf.set_x(pdf.l_margin)  # multi_cell
        if stripped.startswith("## "):
            pdf.set_font("Helvetica", "B", 14)
            pdf.multi_cell(0, 10, stripped[3:])
            pdf.set_font("Helvetica", size=11)
        elif stripped.startswith("# "):
            pdf.set_font("Helvetica", "B", 16)
            pdf.multi_cell(0, 10, stripped[2:])
            pdf.set_font("Helvetica", size=11)
        elif stripped.startswith(("- ", "* ")):
            pdf.multi_cell(0, 8, f"  -  {stripped[2:]}")
        elif stripped:
            pdf.multi_cell(0, 8, stripped)
        else:
            pdf.ln(4)

    if charts:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 10, "Chart Images")
        pdf.set_font("Helvetica", size=11)
        for img_path in _render_chart_images(charts, out_dir):
            pdf.image(img_path, w=180)
            pdf.ln(5)

    path = os.path.join(out_dir, "report.pdf")
    pdf.output(path)
    return path
